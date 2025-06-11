from docx import Document
from datetime import datetime

def generate_dsr_docx(data, output_path):
    doc = Document()

    today_str = datetime.now().strftime("%A, %B %d, %Y")
    doc.add_paragraph(f"Good morning, team,\n\nPlease find below the daily status update for {today_str}.\n")

    for ship in data['ships']:
        doc.add_heading(f"Vessel: {ship['name']}", level=1)
        doc.add_paragraph(f"Region: {ship['region']}")
        doc.add_paragraph("")

        doc.add_heading("T52A TIU – SA 95861", level=2)
        doc.add_paragraph(f"Current Production: {ship['percent_t52a']}")
        doc.add_heading("Work Completed", level=3)
        for item in ship['work_completed_t52a']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading("Planned Work (Next Working Day)", level=3)
        for item in ship['planned_work_t52a']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_paragraph("")
        doc.add_heading("ECS – SA 95676", level=2)
        doc.add_paragraph(f"Current Production: {ship['percent_ecs']}")
        doc.add_heading("Work Completed", level=3)
        for item in ship['work_completed_ecs']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading("Next Steps", level=3)
        for item in ship['next_steps_ecs']:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_paragraph("")

    doc.add_paragraph("Best regards,")
    doc.add_paragraph("OISC John Hughes")
    doc.save(output_path)

    return output_path
